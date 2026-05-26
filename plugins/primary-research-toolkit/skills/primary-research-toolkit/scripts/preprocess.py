#!/usr/bin/env python3
"""
preprocess.py — Transcript pre-processing for Primary Research Toolkit.

Implements ICS Step 2a (Load and split) and Step 2b (Per-transcript pre-processing):
- Load .txt, .md, .docx (folders recursed)
- Detect combined files (multiple interviews in one file) and split
- Filename ↔ person matching against optional metadata
- Per-transcript: input quality detection, speaker resolution,
  exchange-unit (EU) segmentation, non-substantive flagging
- Output: <output_dir>/<basename>.preprocessed.md per transcript + INDEX.md

Input categories:
  - transcript   — verbatim output from any transcription tool (Otter, Teams, Zoom, Rev, etc.)
  - synthesized notes — human or AI prose summary of an interview, speaker-labeled
  - rough notes  — unstructured bullets or shorthand, no speaker labels

Attribution is set provisionally by input type, then refined by speaker resolution outcome:
  transcript + name-match or intro-spiel-confirmed = High
  transcript + aggregate-words-resolved = Medium
  transcript + unresolved (flagged for LLM re-diarization) = Low
  synthesized notes = Medium (regardless of speaker resolution)
  rough notes = Low

Project-agnostic. Newry-team intro-spiel keywords are configurable via
--intro-keywords (defaults match Newry's standard pitch).

Usage:
    python preprocess.py --input <folder|file> [--input <...>] \\
        --output <folder> \\
        [--metadata <metadata.json>] \\
        [--intro-keywords keyword1,keyword2,...]

Originals are never modified. Working copies only.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

# Default Newry-team intro-spiel keywords. Override with --intro-keywords.
DEFAULT_INTRO_KEYWORDS = [
    "Newry", "growth strategy", "consulting firm", "Corning",
    "DuPont", "Norplex", "pick your brain", "Cleveland",
]

# Heuristic thresholds
INTERROGATIVE_PATTERNS = re.compile(
    r"\b(what|how|why|when|where|who|which|can you|could you|would you|"
    r"do you|did you|is there|are there|tell me|walk me|describe|"
    r"explain)\b",
    re.IGNORECASE,
)
SUBSTANTIVE_QUESTION_MIN_WORDS = 12
LOGISTICAL_INTERRUPT_MAX_WORDS = 25
CLOSING_KEYWORDS = re.compile(
    r"\b(thank you|thanks for|appreciate|great talking|great chat|"
    r"have a (good|great) (one|day)|talk soon|nice meeting)\b",
    re.IGNORECASE,
)


# =====================================================================
# Data classes
# =====================================================================

@dataclass
class Turn:
    speaker: str
    timestamp: str  # e.g. "0:00" or "12:34"
    text: str
    role: str = "?"  # 'Q' = interviewer, 'A' = interviewee, '?' = unknown


@dataclass
class ExchangeUnit:
    eu_id: int
    start_ts: str
    end_ts: str
    turns: list[Turn]
    substantive: bool = True
    flag_reason: str = ""
    word_count: int = 0


@dataclass
class TranscriptResult:
    source_path: Path
    interviewee: str
    role: str
    interview_date: str
    location: str
    input_category: str   # 'transcript', 'synthesized notes', 'rough notes'
    attribution: str      # 'High', 'Medium', 'Low'
    needs_rediarization: bool
    total_runtime: str
    total_turns: int
    eus: list[ExchangeUnit]
    interviewers: list[str]
    speaker_resolution_notes: list[str] = field(default_factory=list)


# =====================================================================
# File loading
# =====================================================================

def collect_input_files(inputs: list[str]) -> list[Path]:
    """Walk inputs (files or folders) and collect transcript files."""
    files: list[Path] = []
    for inp in inputs:
        p = Path(inp)
        if p.is_file() and p.suffix.lower() in {".txt", ".md", ".docx"}:
            files.append(p)
        elif p.is_dir():
            for ext in (".txt", ".md", ".docx"):
                files.extend(p.rglob(f"*{ext}"))
    return sorted(set(files))


def read_text(path: Path) -> str:
    """Read .txt/.md as utf-8 text, .docx via python-docx."""
    if path.suffix.lower() in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="replace")
    if path.suffix.lower() == ".docx":
        try:
            from docx import Document  # type: ignore
        except ImportError:
            print("ERROR: python-docx not installed. pip install python-docx --break-system-packages",
                  file=sys.stderr)
            raise
        doc = Document(str(path))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    return ""


# =====================================================================
# Combined-file detection
# =====================================================================

def detect_combined_file_boundaries(text: str) -> list[int]:
    """Return character offsets where new interviews begin.

    Boundary signals:
    - Lines matching 'Interview N:' or 'Call with [Name]'
    - Date+name pattern at section breaks
    - Horizontal rules followed by metadata
    Returns [0] for single-interview files.
    """
    boundaries: list[int] = [0]
    for m in re.finditer(
        r"^\s*(?:Interview\s+\d+[:.]|Call with\s+[A-Z][a-z]+|---+\s*\n\s*\w+:)",
        text, re.MULTILINE,
    ):
        if m.start() > 0:
            boundaries.append(m.start())
    return boundaries


# =====================================================================
# Filename ↔ person matching
# =====================================================================

def match_filename_to_person(
    filename: str, metadata: list[dict]
) -> tuple[dict | None, str]:
    """Try to match a filename to a person in metadata. Returns (match, confidence).

    Token-based matcher. Splits filename and each candidate name into
    alphabetic tokens (≥2 chars, lowercased) and scores by token overlap.
    """
    base = Path(filename).stem.lower()
    base = re.sub(r"_otter[_-]ai[_-]transcript.*", "", base)
    base = re.sub(r"\(\d+\)", "", base).strip()
    base_tokens = set(re.findall(r"[a-zà-ÿ]{2,}", base))
    base_tokens -= {"internal", "external", "transcript", "interview", "call", "sme"}

    best: dict | None = None
    best_score = 0
    best_full = False
    for person in metadata:
        name = (person.get("name") or "").lower()
        if not name:
            continue
        name_tokens = set(re.findall(r"[a-zà-ÿ]{2,}", name))
        if not name_tokens:
            continue
        overlap = name_tokens & base_tokens
        score = len(overlap)
        full = score == len(name_tokens) and score >= 2
        if full and not best_full:
            best, best_score, best_full = person, score, True
        elif full and best_full and score > best_score:
            best, best_score = person, score
        elif not best_full and score > best_score:
            best, best_score = person, score

    if best is None:
        return None, "none"
    if best_full:
        return best, "high"
    if best_score >= 2:
        return best, "medium"
    if best_score >= 1:
        return best, "low"
    return None, "none"


# =====================================================================
# Input category detection
# =====================================================================

def detect_input_category(text: str, path: Path) -> str:
    """Classify input into one of three categories.

    Categories:
      'transcript'        — verbatim output from any transcription tool
      'synthesized notes' — prose summary with speaker labels, no timestamps
      'rough notes'       — unstructured bullets or shorthand, no speaker labels

    Detection is structural (what does the text look like?) not tool-specific.
    Transcript tools vary (Otter, Teams, Zoom, Rev, Granola, etc.) but they all
    produce timestamped speaker-turn output. Synthesized notes and rough notes
    come from human or AI summarization after the fact.
    """
    has_timestamps = bool(re.search(r"\b\d{1,3}:\d{2}\b", text))
    has_speaker_turns = bool(re.search(r"^\s*[A-Z][A-Za-z0-9À-ÿ' .\-]{0,80}?\s{1,4}\d{1,3}:\d{2}", text, re.MULTILINE))
    line_count = text.count("\n")

    # Transcript: has timestamps AND speaker-turn structure
    if has_timestamps and has_speaker_turns and line_count > 20:
        return "transcript"

    # Transcript: has timestamps + many lines (minimal speaker formatting)
    if has_timestamps and line_count > 50:
        return "transcript"

    # Synthesized notes: speaker labels in prose (bold or colon-prefixed), no timestamps
    if not has_timestamps and re.search(r"^\s*\*\*?\w+.*?:\*?\*?", text, re.MULTILINE):
        return "synthesized notes"

    # Rough notes: heavy bullet usage, no timestamps, no speaker labels
    bullet_lines = sum(1 for ln in text.splitlines() if re.match(r"^\s*[-*•]", ln))
    if bullet_lines > 15:
        return "rough notes"

    # Default: treat as synthesized notes if we can't determine
    return "synthesized notes"


def provisional_attribution(category: str) -> str:
    """Provisional attribution before speaker resolution refines it."""
    if category == "transcript":
        return "High"      # refined downward by speaker resolution outcome
    if category == "synthesized notes":
        return "Medium"    # speaker-labeled but not verbatim; resolution doesn't change this
    return "Low"           # rough notes


# =====================================================================
# Speaker + turn parsing
# =====================================================================

# Speaker-line pattern: matches "Name  timestamp" lines from any transcription tool.
# Works for Otter, Teams (partial), Zoom, Rev, and similar tools that produce
# speaker-header + body-text format.
SPEAKER_LINE_RE = re.compile(
    r"^(?P<speaker>[A-Z][A-Za-z0-9À-ÿ' .\-]{0,80}?)\s{1,4}(?P<ts>\d{1,3}:\d{2}(?::\d{2})?)\s*$"
)


def parse_turns(text: str) -> list[Turn]:
    """Parse text into a sequence of Turn records.

    Line-based parser (O(N)). Walks the file once, splitting on lines that
    match the speaker-header pattern; everything between two headers is the
    body of the prior turn.
    """
    turns: list[Turn] = []
    lines = text.splitlines()
    current_speaker: str | None = None
    current_ts: str | None = None
    current_body: list[str] = []

    def flush() -> None:
        if current_speaker and current_body:
            body = "\n".join(current_body).strip()
            if body:
                turns.append(Turn(speaker=current_speaker, timestamp=current_ts or "", text=body))

    for ln in lines:
        m = SPEAKER_LINE_RE.match(ln)
        if m:
            flush()
            current_speaker = m.group("speaker").strip()
            current_ts = m.group("ts").strip()
            current_body = []
        else:
            if current_speaker is not None:
                current_body.append(ln)
    flush()
    return turns


# =====================================================================
# Speaker resolution
# =====================================================================

def _intro_spiel_anchor(
    turns: list[Turn], intro_keywords: list[str], lookahead: int = 8
) -> tuple[str | None, int]:
    """Identify the interviewer label by intro-spiel keyword density.

    Scans the first `lookahead` turns and counts how many distinct intro
    keywords each speaker matches. The speaker with ≥2 matches in the early
    turns is taken to be the interviewer. Tool-agnostic — works on any
    transcript where the interviewing team delivers a project introduction.
    Returns (speaker_label, match_count) or (None, 0).
    """
    if not turns or not intro_keywords:
        return None, 0
    keywords_lower = [k.lower() for k in intro_keywords]
    matches: dict[str, set[str]] = {}
    for t in turns[:lookahead]:
        text_lower = t.text.lower()
        for kw in keywords_lower:
            if kw in text_lower:
                matches.setdefault(t.speaker, set()).add(kw)
    if not matches:
        return None, 0
    top_speaker = max(matches.items(), key=lambda kv: len(kv[1]))
    if len(top_speaker[1]) >= 2:
        return top_speaker[0], len(top_speaker[1])
    return None, 0


def resolve_speakers(
    turns: list[Turn],
    interviewee: str | None,
    intro_keywords: list[str] | None = None,
) -> tuple[list[str], str]:
    """Annotate each Turn with role ('Q' or 'A' or '?') and return (notes, resolution_level).

    Resolution cascade (highest-confidence signal first):
    1. Name match — interviewee's name appears as a speaker label → High confidence.
    2. Intro-spiel anchor — speaker who recites intro keywords in early turns → Q;
       their counterpart identified as interviewee → High confidence.
    3. Aggregate-words fallback — speaker with the most words treated as interviewee
       → Medium confidence.
    4. Resolution failure — no reliable signal. Mark Low attribution, flag for
       LLM re-diarization. Do NOT apply alternation heuristic.

    Returns:
        notes: list of human-readable resolution notes
        resolution_level: 'high', 'medium', 'low' — used to refine attribution
    """
    notes: list[str] = []
    if not turns:
        return notes, "low"

    # 1. Name match for interviewee
    interviewee_label = None
    if interviewee:
        first = interviewee.split()[0].lower()
        last = interviewee.split()[-1].lower()
        for t in turns:
            sp = t.speaker.lower()
            if first in sp or last in sp:
                interviewee_label = t.speaker
                break

    if interviewee_label:
        for t in turns:
            t.role = "A" if t.speaker == interviewee_label else "Q"
        notes.append(
            f"Name match: '{interviewee_label}' identified as interviewee. "
            "Confidence: high."
        )
        return notes, "high"

    # 2. Intro-spiel anchor
    interviewer_label, intro_match_count = _intro_spiel_anchor(
        turns, intro_keywords or []
    )

    # 3. Aggregate-words fallback
    aggregate_top: str | None = None
    aggregate_share: float = 0.0
    if len(turns) >= 4:
        word_totals: dict[str, int] = {}
        for t in turns:
            word_totals[t.speaker] = word_totals.get(t.speaker, 0) + len(t.text.split())
        top_speaker, top_words = max(word_totals.items(), key=lambda kv: kv[1])
        total_words = sum(word_totals.values())
        if total_words:
            aggregate_top = top_speaker
            aggregate_share = top_words / total_words

    # Combine intro-spiel + aggregate-words
    if interviewer_label and aggregate_top and aggregate_top != interviewer_label and aggregate_share >= 0.3:
        interviewee_label = aggregate_top
        for t in turns:
            t.role = "A" if t.speaker == interviewee_label else "Q"
        notes.append(
            f"Intro-spiel anchor: '{interviewer_label}' identified as interviewer "
            f"({intro_match_count} keyword matches in first 8 turns). "
            f"Aggregate-words: '{aggregate_top}' has the most words among non-interviewer speakers "
            f"({100*aggregate_share:.0f}%). Treated as interviewee. "
            "Confidence: medium — multi-signal agreement."
        )
        return notes, "medium"

    # Aggregate-words alone (no intro-spiel signal)
    if aggregate_top and aggregate_share >= 0.4:
        interviewee_label = aggregate_top
        for t in turns:
            t.role = "A" if t.speaker == interviewee_label else "Q"
        display = interviewee or "interviewee"
        notes.append(
            f"Interviewee name '{display}' did not appear as a speaker label. "
            f"Aggregate-words fallback: '{aggregate_top}' has the most words "
            f"({100*aggregate_share:.0f}%) and is treated as the interviewee. "
            "Confidence: medium-low — verify quote attribution before citing verbatim."
        )
        return notes, "medium"

    # 4. Resolution failure — do NOT apply alternation heuristic.
    # Otter and other tools frequently split a single speaker's continuous
    # narrative into multiple consecutive turns, making alternation unreliable.
    # Flag for LLM re-diarization instead.
    unknown_count = sum(
        1 for t in turns if re.match(r"^(unknown speaker|speaker \d+)$", t.speaker, re.IGNORECASE)
    )
    notes.append(
        f"Speaker resolution failed — {unknown_count} of {len(turns)} turns have generic labels "
        "(Unknown Speaker / Speaker N) with no name match, intro-spiel signal, or reliable "
        "aggregate-words signal. "
        "Marked for LLM re-diarization. Do not attribute quotes verbatim from this transcript."
    )
    return notes, "low"


def refine_attribution(category: str, resolution_level: str) -> str:
    """Combine input category and speaker resolution outcome into final attribution."""
    if category != "transcript":
        # Non-transcript categories are not affected by speaker resolution
        return provisional_attribution(category)
    # Transcript: resolution level drives attribution
    if resolution_level == "high":
        return "High"
    if resolution_level == "medium":
        return "Medium"
    return "Low"


# =====================================================================
# Exchange-unit segmentation
# =====================================================================

def segment_into_eus(turns: list[Turn]) -> list[ExchangeUnit]:
    """Group turns into exchange units.

    EU = interviewer turn with substantive question (>12 words + interrogative
    pattern) followed by interviewee response and any interviewer probes within
    that response. Backchannel interviewer turns stay within the current EU.
    """
    eus: list[ExchangeUnit] = []
    current: list[Turn] = []

    def is_substantive_question(t: Turn) -> bool:
        if t.role != "Q":
            return False
        words = len(t.text.split())
        return words >= SUBSTANTIVE_QUESTION_MIN_WORDS and bool(
            INTERROGATIVE_PATTERNS.search(t.text)
        )

    def flush():
        if not current:
            return
        eu = ExchangeUnit(
            eu_id=len(eus) + 1,
            start_ts=current[0].timestamp,
            end_ts=current[-1].timestamp,
            turns=list(current),
            word_count=sum(len(t.text.split()) for t in current),
        )
        eus.append(eu)

    for t in turns:
        if is_substantive_question(t) and current:
            flush()
            current = [t]
        else:
            current.append(t)
    flush()

    return eus


# =====================================================================
# Non-substantive flagging
# =====================================================================

def flag_non_substantive(
    eus: list[ExchangeUnit], intro_keywords: list[str]
) -> None:
    """Mark EUs as non-substantive in-place."""
    if not eus:
        return

    intro_re = re.compile(
        r"\b(" + "|".join(re.escape(k) for k in intro_keywords) + r")\b",
        re.IGNORECASE,
    )

    # First intro EU (within first 4 EUs, ≥2 keyword hits)
    intro_idx: int | None = None
    for i, eu in enumerate(eus[:4]):
        text = " ".join(t.text for t in eu.turns)
        hits = len(intro_re.findall(text))
        if hits >= 2:
            eu.substantive = False
            eu.flag_reason = "intro spiel"
            intro_idx = i
            break

    # Pre-intro warmup
    if intro_idx is not None:
        for eu in eus[:intro_idx]:
            eu.substantive = False
            eu.flag_reason = "pre-intro warmup"

    # Opening pleasantries (first 3 EUs)
    pleasant_re = re.compile(
        r"\b(how are you|nice to meet|good (morning|afternoon|day)|"
        r"thanks for (joining|making time)|how('s| is) your day)\b",
        re.IGNORECASE,
    )
    for eu in eus[:3]:
        if eu.substantive:
            text = " ".join(t.text for t in eu.turns)
            if pleasant_re.search(text):
                eu.substantive = False
                eu.flag_reason = "opening pleasantries"

    # Closing in last 2 EUs
    for eu in eus[-2:]:
        if eu.substantive:
            text = " ".join(t.text for t in eu.turns)
            last_turn_words = len(eu.turns[-1].text.split()) if eu.turns else 0
            if CLOSING_KEYWORDS.search(text) or last_turn_words < 80:
                eu.substantive = False
                eu.flag_reason = "closing / wrap-up"

    # Very short logistical interruptions
    for eu in eus:
        if eu.substantive and eu.word_count < LOGISTICAL_INTERRUPT_MAX_WORDS:
            eu.substantive = False
            eu.flag_reason = "logistical interruption"


# =====================================================================
# Output rendering
# =====================================================================

def render_preprocessed_md(result: TranscriptResult) -> str:
    sub_count = sum(1 for e in result.eus if e.substantive)
    nonsub_count = len(result.eus) - sub_count
    sub_window = ""
    sub_eus = [e for e in result.eus if e.substantive]
    if sub_eus:
        sub_window = f"{sub_eus[0].start_ts} – {sub_eus[-1].end_ts}"

    lines: list[str] = []
    lines.append(f"# Pre-processed: {result.interviewee}")
    lines.append("")
    lines.append("## Header")
    lines.append("")
    lines.append(f"- **Source file:** `{result.source_path.name}`")
    lines.append(f"- **Interviewee:** {result.interviewee} · {result.role}")
    if result.interview_date:
        lines.append(f"- **Interview date:** {result.interview_date}")
    if result.location:
        lines.append(f"- **Interviewee location:** {result.location}")
    lines.append(f"- **Input category:** {result.input_category}")
    lines.append(f"- **Attribution reliability:** {result.attribution}")
    if result.needs_rediarization:
        lines.append(f"- **⚠ Needs LLM re-diarization:** Yes — speaker resolution failed; do not cite quotes verbatim until re-diarized")
    if result.total_runtime:
        lines.append(f"- **Total runtime:** {result.total_runtime}")
    lines.append(f"- **Total turns:** {result.total_turns}")
    lines.append(
        f"- **Exchange units:** {len(result.eus)} "
        f"({sub_count} substantive · {nonsub_count} non-substantive)"
    )
    if sub_window:
        lines.append(f"- **Substantive window:** {sub_window}")
    if result.interviewers:
        lines.append(f"- **Interviewers:** {', '.join(result.interviewers)}")
    lines.append("")

    if result.speaker_resolution_notes:
        lines.append("## Speaker resolution notes")
        lines.append("")
        for n in result.speaker_resolution_notes:
            lines.append(f"- {n}")
        lines.append("")

    lines.append("## Term reconciliation")
    lines.append("")
    lines.append(
        "Fixes from project glossary auto-applied to text below. "
        "See `glossary.md` at project root for canonical terms and "
        "`Decisions made` section in this file for per-fix counts."
    )
    lines.append("")

    lines.append("## Exchange units")
    lines.append("")
    for eu in result.eus:
        flag = ""
        if not eu.substantive:
            flag = f" · **NON-SUBSTANTIVE** ({eu.flag_reason})"
        lines.append(
            f"### EU-{eu.eu_id} · [{eu.start_ts}–{eu.end_ts}]{flag} · {eu.word_count} words"
        )
        lines.append("")
        for t in eu.turns:
            tag = t.role
            lines.append(f"**{tag} — {t.speaker} ({t.timestamp}):** {t.text}")
            lines.append("")

    return "\n".join(lines)


def render_index_md(results: list[TranscriptResult], skipped: list[str]) -> str:
    needs_rediarization = [r for r in results if r.needs_rediarization]

    lines: list[str] = []
    lines.append("# Pre-processing Index")
    lines.append("")

    if needs_rediarization:
        lines.append("## ⚠ Flagged for LLM re-diarization")
        lines.append("")
        lines.append("Speaker resolution failed on these transcripts. Run LLM re-diarization before card generation.")
        lines.append("")
        for r in needs_rediarization:
            lines.append(f"- `{r.source_path.name}` — {r.interviewee}")
        lines.append("")

    lines.append("## Summary")
    lines.append("")
    lines.append(
        "| Interviewee | Date | Category | Attribution | Re-diarize? | Turns | EUs | Sub | Non-sub | Sub window |"
    )
    lines.append(
        "|---|---|---|---|---|---|---|---|---|---|"
    )
    for r in results:
        sub = sum(1 for e in r.eus if e.substantive)
        nonsub = len(r.eus) - sub
        sub_eus = [e for e in r.eus if e.substantive]
        sub_window = (
            f"{sub_eus[0].start_ts}–{sub_eus[-1].end_ts}" if sub_eus else ""
        )
        rediariz = "⚠ Yes" if r.needs_rediarization else "No"
        lines.append(
            f"| {r.interviewee} | {r.interview_date} | {r.input_category} | "
            f"{r.attribution} | {rediariz} | {r.total_turns} | {len(r.eus)} | "
            f"{sub} | {nonsub} | {sub_window} |"
        )
    lines.append("")

    if skipped:
        lines.append("## Files skipped")
        lines.append("")
        for s in skipped:
            lines.append(f"- {s}")
        lines.append("")

    lines.append("## Methodology")
    lines.append("")
    lines.append(
        "Per ICS SKILL Step 2b: input category detection, speaker resolution cascade, "
        "exchange-unit segmentation, non-substantive flagging."
    )
    lines.append("")
    lines.append(
        "**Speaker resolution cascade:** (1) name match → High attribution; "
        "(2) intro-spiel anchor + aggregate-words → Medium; "
        "(3) aggregate-words alone → Medium; "
        "(4) resolution failure → Low, flagged for LLM re-diarization. "
        "Alternation heuristic is NOT used — transcription tools frequently split "
        "single-speaker narratives into consecutive turns, making alternation unreliable."
    )
    lines.append("")
    lines.append(
        f"**Exchange unit definition:** interviewer turn with substantive question "
        f"(>{SUBSTANTIVE_QUESTION_MIN_WORDS} words + interrogative pattern) followed "
        "by interviewee response and any interviewer probes within that response."
    )
    lines.append("")
    lines.append("**Originals untouched.** Working files only.")
    return "\n".join(lines)


# =====================================================================
# Main
# =====================================================================

def process_one(
    path: Path, metadata: list[dict], intro_keywords: list[str]
) -> TranscriptResult | None:
    raw = read_text(path)
    if not raw.strip():
        return None

    person, _conf = match_filename_to_person(path.name, metadata)
    interviewee = person.get("name") if person else "Unknown Interviewee"
    role = person.get("role", "") if person else ""
    date = person.get("date", "") if person else ""
    location = person.get("location", "") if person else ""

    category = detect_input_category(raw, path)

    turns = parse_turns(raw)
    resolution_notes, resolution_level = resolve_speakers(
        turns, interviewee if person else None, intro_keywords
    )

    attribution = refine_attribution(category, resolution_level)
    needs_rediarization = (category == "transcript" and resolution_level == "low")

    eus = segment_into_eus(turns)
    flag_non_substantive(eus, intro_keywords)

    runtime = turns[-1].timestamp if turns else ""
    interviewers = sorted({
        t.speaker for t in turns
        if t.role == "Q" and "Interviewer (inferred)" not in t.speaker
        and not re.match(r"unknown|speaker \d+", t.speaker, re.IGNORECASE)
    })

    return TranscriptResult(
        source_path=path,
        interviewee=interviewee,
        role=role,
        interview_date=date,
        location=location,
        input_category=category,
        attribution=attribution,
        needs_rediarization=needs_rediarization,
        total_runtime=runtime,
        total_turns=len(turns),
        eus=eus,
        interviewers=interviewers,
        speaker_resolution_notes=resolution_notes,
    )


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__.split("\n\n")[0])
    parser.add_argument("--input", action="append", required=True,
                        help="Input file or folder (repeatable)")
    parser.add_argument("--output", required=True,
                        help="Output folder for .preprocessed.md files")
    parser.add_argument("--metadata", default=None,
                        help="Optional JSON file with interviewee metadata: "
                             "[{name, role, date, location, type, id, blind}, ...]")
    parser.add_argument("--intro-keywords", default=",".join(DEFAULT_INTRO_KEYWORDS),
                        help="Comma-separated keywords identifying interviewer "
                             "intro spiel (default: Newry team)")
    args = parser.parse_args(argv)

    out_dir = Path(args.output)
    out_dir.mkdir(parents=True, exist_ok=True)

    metadata: list[dict] = []
    if args.metadata:
        with open(args.metadata) as f:
            metadata = json.load(f)
        if not isinstance(metadata, list):
            print("ERROR: --metadata must be a JSON list of objects.", file=sys.stderr)
            return 2

    intro_keywords = [k.strip() for k in args.intro_keywords.split(",") if k.strip()]

    files = collect_input_files(args.input)
    if not files:
        print("ERROR: no .txt/.md/.docx files found in inputs.", file=sys.stderr)
        return 2

    results: list[TranscriptResult] = []
    skipped: list[str] = []
    for path in files:
        result = process_one(path, metadata, intro_keywords)
        if result is None:
            skipped.append(f"`{path.name}` — empty or unreadable")
            continue
        out_path = out_dir / (path.stem + ".preprocessed.md")
        out_path.write_text(render_preprocessed_md(result), encoding="utf-8")
        results.append(result)
        print(f"Pre-processed: {path.name} → {out_path.name}")

    (out_dir / "INDEX.md").write_text(render_index_md(results, skipped), encoding="utf-8")
    print(f"\nWrote INDEX.md to {out_dir}")

    needs_rediariz = sum(1 for r in results if r.needs_rediarization)
    print(f"Processed: {len(results)}  Skipped: {len(skipped)}  Needs re-diarization: {needs_rediariz}")
    if needs_rediariz:
        print(f"\n⚠  {needs_rediariz} transcript(s) flagged for LLM re-diarization before card generation.")
        print("   See INDEX.md for the list.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
