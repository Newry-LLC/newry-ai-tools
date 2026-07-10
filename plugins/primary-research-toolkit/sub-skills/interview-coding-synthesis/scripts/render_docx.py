#!/usr/bin/env python3
"""
render_docx.py — Deterministic Markdown -> Word renderer for ICS cards/Roll-up docs.

Why this exists: ICS cards and Roll-ups were previously produced by converting
Markdown through pandoc. Pandoc's presence on a consultant machine is not
guaranteed (no install step ships with this plugin), and a prior update run
produced a different style set than the first run plus leaked literal ">"
characters into body text. This script replaces that step with a single,
dependency-pinned (python-docx only) renderer that always emits the same
styles from the same constrained Markdown dialect ICS actually produces:
headers (#-#####), GFM pipe tables, bullet lists (2-space nested indent),
bold (**text**), horizontal rules (---), plain paragraphs.

The canonical Markdown file is always the source of truth. This script always
FULLY REBUILDS the docx from that Markdown — it never edits an existing docx
in place and never reads a docx back to append to it.

Usage:
    python render_docx.py --input <canonical.md> --output <Output.docx>

Dependencies: python-docx (pip install python-docx --break-system-packages)
Chain with style_docx.py afterward for coverage-table cell shading:
    python render_docx.py --input cards.md --output "Cards v3.docx"
    python style_docx.py --input "Cards v3.docx" --output "Cards v3.docx"
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path


def _add_bold_runs(paragraph, text: str) -> None:
    """Split on **bold** markers and add runs, preserving bold spans."""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)


def _table_line(line: str) -> list[str] | None:
    stripped = line.strip()
    if not stripped.startswith("|"):
        return None
    cells = [c.strip() for c in stripped.strip("|").split("|")]
    return cells


def _is_table_separator(cells: list[str]) -> bool:
    # GFM separator row: cells like ---, :---, ---:, :---:
    return all(re.fullmatch(r":?-{2,}:?", c) for c in cells)


def render(md_text: str, doc) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore

    lines = md_text.splitlines()
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.rstrip()

        if not stripped.strip():
            i += 1
            continue

        # Horizontal rule
        if re.fullmatch(r"-{3,}", stripped.strip()):
            doc.add_paragraph().add_run("")  # spacer; no visible rule glyph needed
            i += 1
            continue

        # Headers
        header_match = re.match(r"^(#{1,5})\s+(.*)$", stripped)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2).strip()
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Blockquote lines are NOT part of the ICS card/Roll-up dialect.
        # If one appears (a template regression), strip the marker rather
        # than emit a literal ">" into the document.
        if stripped.lstrip().startswith(">"):
            text = re.sub(r"^\s*>\s?", "", stripped)
            p = doc.add_paragraph()
            _add_bold_runs(p, text)
            i += 1
            continue

        # Tables (GFM pipe tables): header row, separator row, body rows
        cells = _table_line(stripped)
        if cells is not None:
            table_rows = [cells]
            j = i + 1
            saw_separator = False
            while j < n:
                next_cells = _table_line(lines[j].rstrip())
                if next_cells is None:
                    break
                if not saw_separator and _is_table_separator(next_cells):
                    saw_separator = True
                    j += 1
                    continue
                table_rows.append(next_cells)
                j += 1

            n_cols = max(len(r) for r in table_rows)
            table = doc.add_table(rows=len(table_rows), cols=n_cols)
            table.style = "Table Grid"
            for r_idx, row in enumerate(table_rows):
                for c_idx in range(n_cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    cell_para = table.cell(r_idx, c_idx).paragraphs[0]
                    _add_bold_runs(cell_para, text)
                    if r_idx == 0:
                        for run in cell_para.runs:
                            run.bold = True
            i = j
            continue

        # Bullet lists (nested via 2-space indent per level)
        bullet_match = re.match(r"^(\s*)([-*])\s+(.*)$", line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            level = indent // 2
            text = bullet_match.group(3).strip()
            style = "List Bullet" if level == 0 else f"List Bullet {min(level + 1, 3)}"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
            _add_bold_runs(p, text)
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        _add_bold_runs(p, stripped.strip())
        i += 1


def render_docx(md_path: Path, out_path: Path) -> int:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        print(
            "ERROR: python-docx not installed. "
            "pip install python-docx --break-system-packages",
            file=sys.stderr,
        )
        return 2

    if not md_path.is_file():
        print(f"ERROR: input Markdown not found: {md_path}", file=sys.stderr)
        return 2

    md_text = md_path.read_text(encoding="utf-8")
    doc = Document()
    render(md_text, doc)
    doc.save(str(out_path))
    print(f"Rendered {md_path.name} -> {out_path}")
    return 0


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__.split("\n\n")[0])
    parser.add_argument("--input", required=True, help="Canonical Markdown input path")
    parser.add_argument("--output", required=True, help="Output docx path (always fully rewritten)")
    args = parser.parse_args(argv)

    return render_docx(Path(args.input), Path(args.output))


if __name__ == "__main__":
    sys.exit(main())
