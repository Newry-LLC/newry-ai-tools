# -*- coding: utf-8 -*-
import re, json
from docx import Document
from docx.shared import Pt, RGBColor
def wc(s): return len(re.findall(r"\S+", s))

# ---- firm-level info from corrected JSON ----
_corr = json.load(open("skills/sequoia-workflow/cache/Walnut_Creek-firms-2026-07-03-corrected.json",encoding="utf-8"))["firms"]
def _norm(s): return re.sub(r'[^a-z0-9]','',(s or '').lower())
def _core(s):  # norm + strip trailing corp suffixes
    n=_norm(s)
    for suf in ("incorporated","inc","llc","lp","llp","lllp","ltd","corp","co","company","sf"):
        if n.endswith(suf): n=n[:-len(suf)]
    return n
_byname = {f["name"]: f for f in _corr}
_bynorm = {_norm(f["name"]): f for f in _corr}
_bycore = {_core(f["name"]): f for f in _corr}
def _rec(fn):
    return (_byname.get(fn) or _bynorm.get(_norm(fn)) or _bynorm.get(_norm(fn.split(",")[0]))
            or _bycore.get(_core(fn)) or _bycore.get(_core(fn.split(",")[0])))
def _aum(v):
    try: n=float(v)
    except: return None
    if n<=0: return None
    return f"${n/1e9:.1f}B" if n>=1e9 else f"${n/1e6:.0f}M"
def _clean(s):
    s=re.sub(r'\s+',' ',str(s)).strip()
    s=re.sub(r'^\[[^\]]*\]\s*','',s)   # strip leading meta like "[Web research 2026-07-03 ...]"
    return s
def firm_info(fn):
    c=_rec(fn)
    if not c: return []
    rows=[]
    aum=_aum(c.get("aum"))
    emp=c.get("employees")
    line1=[]
    if aum: line1.append(f"AUM: {aum}")
    if emp: line1.append(f"Team: {emp}")
    if c.get("fee_structure"): line1.append(f"Fees: {_clean(c['fee_structure'])}")
    if line1: rows.append(" · ".join(line1))
    if c.get("service_model"): rows.append(f"Service model: {_clean(c['service_model'])}")
    if c.get("typical_client"): rows.append(f"Typical client: {_clean(c['typical_client'])}")
    if c.get("platform_technology"): rows.append(f"Platform / tech: {_clean(c['platform_technology'])}")
    if c.get("awards"): rows.append(f"Awards / recognition: {_clean(c['awards'])}")
    if c.get("other_firm_notes"): rows.append(f"Notes: {_clean(c['other_firm_notes'])}")
    return rows

OPENING = "I wanted to follow up on my earlier note. As I mentioned, Sequoia is looking to expand in the Bay Area, and I thought it might help to share a bit more about how they approach partnership."
WRAPPER = "Sequoia is deliberate about the number of partnerships they pursue, focusing on cultural and service-model alignment above all else. This has supported ~98% client and advisor retention while generating 18–20% organic growth net of market. They take their time pacing integrations, keep client service models intact, and keep founders meaningfully involved in what they've built."
ASK = "If it would be useful, I'd welcome a brief, low-pressure conversation to share how this works in practice and hear how you're thinking about the firm's next chapter. If now isn't the right time, I completely understand."

# (firm_name, subject_short, greeting, middle_paragraph)
EMAILS = [
("Private Wealth Partners, LLC","Private Wealth Partners","Bill",
 "In looking at Private Wealth Partners, what stood out to me was its depth of tenure and long-term orientation — more than three decades building the firm and its predecessor, with a concentrated, relationship-driven approach to UHNW families and charitable organizations. That kind of continuity is exactly what Sequoia looks for in a partner: a firm worth preserving, where the right partnership removes friction behind the scenes rather than changing what already works."),
("Intellectus Partners, LLC","Intellectus Partners","David",
 "What caught my eye about Intellectus was its intentional focus on entrepreneurs and innovators whose wealth is tied to their business interests — a platform built to engage clients at every stage of wealth creation, not just after liquidity. That clarity of focus is exactly what Sequoia values in a partner: their approach supports what you've built rather than changing it, adding depth behind the scenes while your client relationships stay in your hands."),
("Mission Creek Capital Partners, Inc.","Mission Creek Capital","Henri",
 "What stood out to me about Mission Creek was your specialized focus on concentrated equity positions and post-venture transition planning — complex balance-sheet work that takes both technical depth and real trust, and that your sub-advisory relationships with other advisors and family offices clearly reflect. That kind of specialized capability is exactly what Sequoia looks to preserve in a partner, adding infrastructure behind the scenes while your investment process and client relationships stay intact."),
("Blume Capital Management, Inc.","Blume Capital Management","Jeffrey",
 "I was struck by Blume Capital's longevity and research-driven investment philosophy — three decades managing portfolios for individuals, families, and foundations with a value-oriented, non-institutional approach and genuinely personalized counsel. That depth of client relationship is exactly what Sequoia seeks in a partner: their approach keeps what you do best in your hands, adding operational support behind the scenes rather than changing what's working."),
("Stewart Wealth Management, Inc.","Stewart Wealth Management","Ben",
 "When I read about Stewart Wealth Management, I was struck by the straightforward, client-first culture you've built since founding the firm in 2007 — a fee-only, planning-led model for Marin and Bay Area families centered on long-term partnership rather than product. Firms built around one clear vision and sustained over time are exactly what Sequoia looks for, and their approach protects that, keeping your client relationships and service model intact while adding support behind the scenes."),
("Ohana Advisors","Ohana Advisors","Ohana Team",
 "In looking at Ohana, what stayed with me was your referral-only model and 25-year track record serving ultra-high-net-worth families with genuine multi-family-office depth — manager selection, direct investments, estate planning, family-office administration, and philanthropic advisory, all in a boutique setting. That kind of relationship-built practice is exactly what Sequoia looks to preserve, adding infrastructure behind the scenes while your client relationships and boutique feel stay intact."),
("Roof Eidam Maycock Peralta LLC","Roof Eidam Maycock Peralta","Gary",
 "I was drawn to Roof Eidam Maycock Peralta for your multi-decade continuity and life-planning orientation — partners with 25 or more years at the firm, and a referral-only model that balances investment management with comprehensive personal planning. That culture, built around service quality rather than growth for its own sake, is exactly what Sequoia looks to preserve, removing friction behind the scenes rather than changing what already works."),
("DiversiFi Capital LLC","DiversiFi Capital","Nathan",
 "What caught my eye about DiversiFi was your focused specialization in equity-compensation planning for Bay Area tech professionals — RSUs, stock options, AMT planning, and concentrated positions, served through a virtual model that reaches clients wherever their careers take them. That clear positioning in a well-defined segment is exactly what Sequoia values, and their approach supports it, keeping your client relationships in your hands while adding depth behind the scenes."),
("Elevation Wealth Partners, LLC","Elevation Wealth Partners","Barry",
 "When I looked at Elevation Wealth Partners, I was struck by your genuine integration of investment management with in-house tax preparation and accounting — offering CPA services alongside portfolio management rather than outsourcing them. That kind of integrated, client-first platform is exactly what Sequoia looks for in a partner, and their approach keeps what you do best in your hands, adding operational depth behind the scenes rather than changing what's working."),
("Silicon Valley Capital Partners, L.P.","Silicon Valley Capital Partners","Cindy",
 "Silicon Valley Capital Partners stood out to me for your proprietary, research-driven investment process — a top-down macroeconomic framework paired with bottom-up valuation and a momentum discipline that reflects real intellectual rigor. That distinctive process is exactly what Sequoia looks to preserve in a partner, providing infrastructure and support behind the scenes while your investment process and client relationships stay fully intact."),
("Wealthvalues, LLC","Wealthvalues","Karen",
 "What caught my eye about Wealthvalues was the depth of your work with Silicon Valley executives — guiding clients through IPOs, acquisitions, and genuinely complex compensation, delivered as fully personalized, fee-only wealth management. That kind of focused expertise in a well-defined client base is exactly what Sequoia values in a partner, and their approach supports it, keeping your client relationships in your hands while adding depth behind the scenes."),
("Frank, Rimerman Advisors LLC","Frank, Rimerman Advisors","Frank, Rimerman Team",
 "What stood out to me about Frank, Rimerman Advisors is the genuinely integrated platform you've built — an investment practice tied to one of Silicon Valley's largest independent CPA firms, serving UHNW entrepreneurs and executives with tax and wealth advice under one roof. That kind of integrated model is exactly what Sequoia looks to preserve in a partner, adding infrastructure behind the scenes while your client relationships and integrated approach stay intact."),
("Neumann Capital Management","Neumann Capital","Dan",
 "What caught my eye about Neumann Capital was your collaborative team culture and the CFA-led investment discipline at its core — a fee-only, fiduciary approach built around long-standing client relationships. That kind of disciplined, relationship-first practice is exactly what Sequoia looks for in a partner, and their approach keeps what you do best in your hands, adding operational support behind the scenes rather than changing what's working."),
("Wealth Architects, LLC","Wealth Architects","Mark",
 "What stood out to me about Wealth Architects was the breadth of the multi-shareholder team you've built and the purpose-driven approach behind it — integrating financial planning with insurance, tax, estate, and real estate for Silicon Valley executives and families. That kind of comprehensive, team-based platform is exactly what Sequoia looks to preserve in a partner, adding depth behind the scenes while your client relationships stay in your hands."),
("Spiegelman Wealth Management, Inc.","Spiegelman Wealth Management","Adam",
 "I was struck by Spiegelman Wealth's generational story and the deliberate independence behind launching your own RIA in 2025 — taking a multigenerational family practice and building it into an independent, fee-only firm for Walnut Creek families. That kind of conviction and continuity is exactly what Sequoia looks for in a partner, and their approach protects it, keeping your client relationships and service model intact while adding support behind the scenes."),
("Fiduciary Financial Group, LLC","Fiduciary Financial Group","Richard",
 "What caught my eye about Fiduciary Financial Group was the strength of your multi-credential team and the integration of CPA services with investment advisory — in-house tax reporting and planning alongside evidence-based portfolio management. That kind of integrated, client-first platform is exactly what Sequoia looks for in a partner, and their approach keeps what you do best in your hands, adding operational depth behind the scenes rather than changing what's working."),
("Traveka Wealth, LLC","Traveka Wealth","Tony",
 "What stood out to me about Traveka was your clarity and focus — an independent RIA built around a defined niche in equity-compensation planning for Silicon Valley technology professionals, from RSUs and stock options to concentrated positions. That kind of clear positioning in a well-defined segment is exactly what Sequoia values in a partner, and their approach supports it, keeping your client relationships in your hands while adding depth behind the scenes."),
("Elmwood Wealth Management","Elmwood Wealth Management","Bob and Shannon",
 "I was struck by the partnership behind Elmwood — a team that left institutional asset management to serve Bay Area technology professionals on its own terms, with real depth in equity compensation from RSUs to QSBS. That kind of intentional, client-first practice is exactly what Sequoia looks for in a partner, and their approach keeps your client relationships in your hands while adding support behind the scenes."),
("Blueprint Investing LLC","Blueprint Investing","Bob",
 "What caught my eye about Blueprint was its deliberate concentration — a boutique practice built around a small base of HNW business leaders, entrepreneurs, and families navigating rapid wealth creation and complex personal planning. That kind of focused, relationship-first model is exactly what Sequoia looks for in a partner, and their approach keeps what you do best in your hands, adding depth behind the scenes rather than changing what's working."),
("Clarity Wealth Advisors","Clarity Wealth Advisors","Parvin",
 "What stood out to me about Clarity Wealth Advisors was your depth of experience and genuinely client-first philosophy — an independent, fee-only practice sustained for over a decade, with a real education-focused approach to long-term relationships. That kind of intentional culture is exactly what Sequoia looks for in a partner, and their approach protects it, keeping your client relationships and service model intact while adding support behind the scenes."),
("Capital Trust Advisors","Capital Trust Advisors","John",
 "What impressed me about Capital Trust was the depth and tenure of your team — a disciplined investment process spanning commercial lending, fixed income, and portfolio management, brought to a boutique setting with real institutional rigor. That kind of high-touch, multi-generational focus is exactly what Sequoia looks to preserve in a partner, adding infrastructure behind the scenes while your client relationships stay intact."),
("Avalon Capital Management","Avalon Capital Management","Clara and David",
 "I was struck by Avalon's intellectual rigor and the longevity of your founding partnership — deep roots in quantitative research, including your proprietary “Ada” dynamic-allocation model, sustained and refined for over three decades. That kind of systematic discipline and continuity is exactly what Sequoia looks for in a partner, and their approach keeps your investment process and client relationships intact while adding support behind the scenes."),
("Everest Private Wealth","Everest Private Wealth","Ranga and Ramprasad",
 "Everest's founding story is what stayed with me — a firm built specifically to close the gaps in financial services that Silicon Valley professionals so often run into, focused on first-generation high earners with complex, multi-account wealth. That kind of well-defined niche and origin is exactly what Sequoia looks for in a partner, and their approach supports it, keeping your client relationships in your hands while adding depth behind the scenes."),
("Three Bridge Wealth Advisors","Three Bridge Wealth Advisors","Eric and Brett",
 "Three Bridge is on my radar for its serious positioning at the top of the Silicon Valley market — deep expertise in alternative investments, a personal-CFO service model, and national media visibility that reflects the credibility you've earned with entrepreneurs, VCs, and executives. That kind of comprehensive, high-touch platform is exactly what Sequoia looks to preserve in a partner, adding infrastructure behind the scenes while your client relationships stay intact."),
("Opes Wealth Management, LLC","Opes Wealth Management","Mark and Erin",
 "What stood out to me about Opes was the way you integrate real estate advisory into the wealth management relationship — pairing deep investment management with real estate expertise for Apple and Silicon Valley executives weighing equity compensation alongside property decisions. That kind of distinctive, integrated model is exactly what Sequoia looks to preserve in a partner, adding depth behind the scenes while your client relationships stay in your hands."),
("Rosenblum Silverman Sutton SF","Rosenblum Silverman Sutton","John",
 "What caught my eye about Rosenblum Silverman Sutton was your remarkable institutional history — founded in 1984 around independently developed investment disciplines, with more than 40 years of continuity serving high-net-worth families, trusts, and foundations. That legacy of conviction and independence is exactly what Sequoia looks to preserve in a partner, adding infrastructure behind the scenes while your investment process and client relationships stay intact."),
("Chequers Financial Management LLC","Chequers Financial Management","Megan",
 "What stood out to me about Chequers was your clear positioning at the intersection of law, tax, and wealth management — deep professional roots that earn credibility with entrepreneurs and executives navigating genuinely complex tax, liquidity, and estate decisions. Your recognition as one of the fastest-growing RIAs in the country signals real client resonance, and that kind of momentum is exactly what Sequoia looks for in a partner."),
("Atlas Capital Advisors Inc.","Atlas Capital Advisors","Jonathan",
 "I was struck by the institutional pedigree of the Atlas team and the seriousness of your process — deep institutional investment experience, including large-scale corporate treasury and CIO-level portfolio management, applied through a systematic, multi-factor equity strategy. That kind of rigor is rare in independent advisory, and it's exactly what Sequoia looks to preserve in a partner, adding support behind the scenes while your investment process stays intact."),
("Flywheel Private Wealth LLC","Flywheel Private Wealth","Brandon",
 "What makes Flywheel interesting to me is the deliberateness of its founding — a fee-only, fiduciary RIA launched in 2023 on years of client-relationship experience, built as a low-ego, high-touch boutique by design. That kind of intentional, client-first practice is exactly what Sequoia looks for in a partner, and their approach keeps your client relationships in your hands while adding depth behind the scenes."),
("McCarthy Asset Management, Inc.","McCarthy Asset Management","Stephen",
 "What stood out to me about McCarthy Asset Management was your tightly integrated CPA and advisory model — an investment practice launched in 1999 in direct response to what clients were asking for, pairing tax-integrated planning with a disciplined, long-term approach. That kind of client-driven integration is exactly what Sequoia looks for in a partner, and their approach keeps what you do best in your hands while adding support behind the scenes."),
("Brickley Wealth Management","Brickley Wealth Management","Stephen",
 "What caught my eye about Brickley was the integration of your CPA and investment advisory roots and the multi-generational depth you've built with professional families — a practice that began as a tax firm in 1985 and expanded into comprehensive wealth management. That kind of tax-integrated, relationship-first model is exactly what Sequoia looks to preserve in a partner, adding depth behind the scenes while your client relationships stay intact."),
("Jackson Square Capital, LLC","Jackson Square Capital","Andrew",
 "What stood out to me about Jackson Square was your positioning at the ultra-high-net-worth end of the market — deep institutional portfolio-management experience applied to a highly selective client base of senior executives and family offices. That kind of customized, high-touch practice is exactly what Sequoia looks to preserve in a partner, adding infrastructure behind the scenes while your client relationships stay in your hands."),
("Gordian Wealth Advisors, LLC","Gordian Wealth Advisors","Elliott",
 "What caught my eye about Gordian was the distinctive combination behind it — a foundation in institutional markets and analytical consulting brought into an independent advisory practice for high-net-worth families across Northern California. That kind of analytical rigor paired with independence is exactly what Sequoia looks for in a partner, and their approach keeps your client relationships in your hands while adding depth behind the scenes."),
("Asset Dedication, LLC","Asset Dedication","Steve and Brent",
 "I was struck by the intellectual originality behind Asset Dedication — an academic, decision-sciences foundation turned into an investment platform centered on bond laddering and dedicated portfolios for retirement income, serving both direct clients and other advisors. That kind of distinctive, research-driven approach is exactly what Sequoia looks to preserve in a partner, adding infrastructure behind the scenes while your investment process stays intact."),
("Destination Wealth Management","Destination Wealth Management","Michael",
 "What stood out to me about Destination was your longevity and national profile — founded in 1986 as one of the Bay Area's oldest independent RIAs and sustained for nearly four decades, with consistent national recognition and a research-driven investment approach. That kind of durable, well-recognized platform is exactly what Sequoia looks to preserve in a partner, adding support behind the scenes while your client relationships stay intact."),
("Sivia Capital Partners, LLC","Sivia Capital Partners","Benjamin",
 "What caught my eye about Sivia was your positioning at the top of the UHNW market — a multi-family office built for entrepreneurs and multigenerational families across the full startup-to-exit-to-generational wealth cycle, with an exceptionally high average relationship size. That kind of specialized, high-touch practice is exactly what Sequoia looks to preserve in a partner, adding infrastructure behind the scenes while your client relationships stay in your hands."),
("Creekside Partners","Creekside Partners","Richard and Andrew",
 "What stood out to me about Creekside was your distinctive community orientation — a deliberate regional focus in the North Bay since 2008, pairing portfolio management with educational seminars and workshops that reflect a genuine commitment to your community. That kind of relationship-first, community-rooted practice is exactly what Sequoia looks for in a partner, and their approach keeps your client relationships in your hands while adding support behind the scenes."),
("Ridgecrest Wealth Partners, LLC","Ridgecrest Wealth Partners","Steven",
 "What caught my eye about Ridgecrest was the scale you've built since founding in 2021 — over $700M in assets and 744 client relationships from a three-person team in just a few years points to real operational discipline. That kind of efficient, fast-growing model is exactly what Sequoia looks for in a partner, and their approach supports it, keeping your client relationships in your hands while adding depth behind the scenes."),
("Fairview Capital","Fairview Capital","Andrew and Peter",
 "What stood out to me about Fairview was your multi-generational continuity and the depth you've built over nearly three decades of independent practice — seasoned investment experience now carried into the next generation of leadership, serving HNW families and foundations nationally. That kind of continuity is exactly what Sequoia looks to preserve in a partner, adding infrastructure behind the scenes while your client relationships stay intact."),
("Financial Avengers, Inc.","Financial Avengers","Thomas",
 "What caught my eye about Financial Avengers was the focused practice you've built in the East Bay — a lean operation serving individuals and small businesses across more than 200 relationships with a strong, personal, client-first approach. That kind of hands-on, relationship-driven model is exactly what Sequoia looks for in a partner, and their approach keeps your client relationships in your hands while adding support behind the scenes."),
("Capital Advantage, Inc.","Capital Advantage","John",
 "What stood out to me about Capital Advantage was the longevity and scale you've built independently in the East Bay — a practice sustained through multiple market cycles and now managing over $1.25B, one of the larger advisory firms in the Walnut Creek corridor. That kind of durable, at-scale independence is exactly what Sequoia looks to preserve in a partner, adding infrastructure behind the scenes while your client relationships stay intact."),
("Westend Capital Management, LLC","Westend Capital Management","George",
 "What caught my eye about Westend was your growth trajectory and the concentrated client base you've built in Marin — strong net-new-asset and revenue growth alongside a relatively unusual willingness to align fees with performance. That kind of momentum and alignment is exactly what Sequoia looks for in a partner, and their approach supports it, keeping your client relationships in your hands while adding depth behind the scenes."),
("Occidental Asset Management, LLC","Occidental Asset Management","Charles and Nathan",
 "What stood out to me about Occidental was your genuinely differentiated positioning around behavioral finance — operating as “Your Mental Wealth Advisors,” you integrate financial psychology with traditional investment management, tax, and estate planning for a broad client base. That kind of distinctive approach is exactly what Sequoia looks to preserve in a partner, adding infrastructure behind the scenes while your client relationships stay intact."),
("Bell Investment Advisors Inc","Bell Investment Advisors","James",
 "What caught my eye about Bell Investment Advisors was the depth of practice you've built in the East Bay over more than 25 years — nearly $800M across a broad service menu spanning planning, portfolio management, and coaching, serving individuals, families, and institutions. That kind of established, full-service practice is exactly what Sequoia looks for in a partner, and their approach keeps your client relationships in your hands while adding support behind the scenes."),
("Sentry Advisors, LLC","Sentry Advisors","Alex",
 "I was struck by the remarkable concentration and scale of your client relationships at Sentry — a four-person firm managing over $2.4B, with average client assets near $13M, reflecting a highly selective practice built around a small number of substantial relationships. That kind of focused, high-caliber model is exactly what Sequoia looks to preserve in a partner, adding infrastructure behind the scenes while your client relationships stay intact."),
("Greykasell Wealth Strategies, Inc.","Greykasell Wealth Strategies","David",
 "What caught my eye about Greykasell was the scale you brought in at launch — registering in 2024 with over $443M in assets from the outset points to an experienced team that had already built a substantial, loyal client following. That kind of established book and client trust is exactly what Sequoia looks for in a partner, and their approach keeps your client relationships in your hands while adding depth behind the scenes."),
("One Wealth Advisors, LLC","One Wealth Advisors","David and Jonathan",
 "What stood out to me about One Wealth was your distinctive founding philosophy and the breadth you've built since the team left J.P. Morgan together in 2014 — a “Financial Life Planning” model that treats financial and personal wellbeing as inseparable. That kind of clear point of view is exactly what Sequoia looks for in a partner, and their approach protects it, keeping your client relationships and service model intact while adding support behind the scenes."),
("Navigation Group LLC","Navigation Group","Erik",
 "What caught my eye about Navigation Group was the growth you've achieved on the Peninsula since founding in 2017 — over $1B in assets and more than 340 client relationships across a six-advisor team reflects a practice that has built real scale in a short time. That kind of momentum is exactly what Sequoia looks for in a partner, and their approach supports it, keeping your client relationships in your hands while adding depth behind the scenes."),
("Sand Hill Global Advisors, LLC","Sand Hill Global Advisors","Brian",
 "What stood out to me about Sand Hill was your institutional investment infrastructure and deep roots in the Silicon Valley ecosystem — an employee-owned firm pairing open-architecture, globally oriented portfolio management with decades in the local tech and venture community. That kind of institutional depth and continuity is exactly what Sequoia looks to preserve in a partner, adding support behind the scenes while your client relationships stay intact."),
("TTP Investments, Inc.","TTP Investments","Harris",
 "What caught my eye about TTP was the focused practice you've built in the South Bay since founding in 2016 — over $570M across nearly 400 client relationships from a lean team reflects steady, earned trust in the San Jose market. That kind of efficient, relationship-first model is exactly what Sequoia looks for in a partner, and their approach keeps your client relationships in your hands while adding support behind the scenes."),
("RPG Investment Advisory, LLC","RPG Investment Advisory","Matthew",
 "What stood out to me about RPG was the operational scale your team has achieved in the Tri-Valley — nearly $1.05B across more than 950 client relationships reflects both efficiency and breadth, a model built to serve a broad client base well. That kind of scalable, high-volume practice is exactly what Sequoia looks for in a partner, and their approach supports it, keeping your client relationships in your hands while adding depth behind the scenes."),
("Verita Strategic Wealth Partners, LLC","Verita Strategic Wealth Partners","Kelly",
 "What stood out to me about Verita was your positioning at the very top of the ultra-high-net-worth market — an exceptionally selective practice, launched in 2025, serving a small number of substantial family relationships bi-coastally with a genuinely high-touch model. That kind of exclusive, high-caliber focus is exactly what Sequoia looks to preserve in a partner, adding infrastructure behind the scenes while your client relationships stay in your hands."),
("Van Strum & Towne Inc","Van Strum & Towne","Barbara",
 "What stayed with me about Van Strum & Towne was something simple and increasingly rare: longevity. Founded in 1976 and now approaching fifty years, managing nearly $500M for a select group of high-net-worth individuals, families, and institutions, the firm reflects a durability few independents achieve. That kind of continuity is exactly what Sequoia looks to preserve in a partner, adding support behind the scenes while your client relationships stay intact."),
("The Pacific Center For Financial Services","The Pacific Center for Financial Services","Stephen",
 "What caught my eye about The Pacific Center was the depth of practice you've built in the Tri-Valley since founding in 1997 — nearly three decades of independent advisory work managing close to $1B for individuals and small businesses in the San Ramon area. That kind of established, community-rooted practice is exactly what Sequoia looks for in a partner, and their approach keeps your client relationships in your hands while adding support behind the scenes."),
("Vantage Wealth Management, LLC","Vantage Wealth Management","Val",
 "What stood out to me about Vantage was your equal five-way partnership structure and the scale the team has built together on the Peninsula — over $1.1B across more than 600 accounts reflects a genuinely collaborative model rather than a single-founder practice. That kind of shared, durable ownership is exactly what Sequoia looks for in a partner, and their approach keeps your client relationships in your hands while adding depth behind the scenes."),
("North Berkeley Wealth Management, LLC","North Berkeley Wealth Management","Kate",
 "What caught my eye about North Berkeley was the culture you've built and your clear sense of place in the East Bay — managing over $830M with a sixteen-person team, grounded in the Berkeley community. That kind of team depth and local rootedness is exactly what Sequoia looks for in a partner, and their approach keeps your client relationships in your hands while adding support behind the scenes."),
("Vernal Point Advisors, LLC","Vernal Point Advisors","Paul and Mark",
 "What stood out to me about Vernal Point was the rarity of your positioning — a firm purpose-built around single-family-office design and management for ultra-high-net-worth global families and Fortune-400-caliber family offices, serving a small number of exceptionally substantial relationships. That kind of specialized, high-caliber focus is exactly what Sequoia looks to preserve in a partner, adding infrastructure behind the scenes while your client relationships stay in your hands."),
("Westhill Financial Advisors, Inc.","Westhill Financial Advisors","Kirk and Meghan",
 "What caught my eye about Westhill was your 35-year track record and the thoughtful succession you've put in place — a firm dating to 1989, rebranded as Westhill in 2021, now elevating a new generation of leadership while keeping its client-first, community-focused approach intact. That kind of continuity and deliberate planning is exactly what Sequoia looks to preserve in a partner, adding support behind the scenes while your client relationships stay intact."),
("Werba Rubin Papier Wealth Management, LLC","Werba Rubin Papier Wealth Management","Alan, Aaron, and Jason",
 "What stood out to me about Werba Rubin Papier was the strength of your three-partner founding model and the scale you've built together in the South Bay — equal-thirds ownership sustaining a firm now managing over $1.4B with an evidence-based investment approach. That kind of shared ownership and durable scale is exactly what Sequoia looks for in a partner, and their approach keeps your client relationships in your hands while adding support behind the scenes."),
]

FLAGS = [
 "Rosenblum Silverman Sutton — round-1 contact (John Stilwell) flagged as a possible owner-name mismatch; confirm the right recipient before sending.",
 "Neumann Capital — open question on current ownership (“may have sold?”); paragraph kept firm-level and neutral. Confirm before sending.",
 "Sivia Capital — open question on the named contact's current day-to-day involvement; paragraph kept firm-level.",
 "Occidental / Creekside / Westend — founder/co-founder attribution was disputed in earlier verification, so paragraphs deliberately avoid naming a founder and stay firm-level.",
 "Sage Rhino Capital — NOT drafted; parked on the 2018 lawsuit finding, needs a Jack/AJ go/no-go.",
 "Lyell Wealth Management — excluded (declined).",
]

doc = Document()
f=doc.styles['Normal'].font; f.name='Calibri'; f.size=Pt(11)
doc.add_heading("Walnut Creek — Second (Follow-Up) Emails — v2 DRAFT",level=1)
p=doc.add_paragraph("First-person voice, firm-level content, matched to the Walnut Creek initial reach-outs (style + length). Structure: standard opening → firm-level paragraph → Sequoia wrapper → low-pressure ask. Recipients = only those emailed in round 1. All %d firms below. See “Flags for review” at the end." % len(EMAILS))
p.runs[0].italic=True

for i,(fn,subj,greet,mid) in enumerate(EMAILS,1):
    doc.add_heading(f"{i}. {fn}",level=2)
    m=doc.add_paragraph(); r=m.add_run(f"To: {greet}    |    Subject: Sequoia / Eide Bailly: Bay Area expansion and {subj}"); r.bold=True
    doc.add_paragraph(f"{greet},")
    doc.add_paragraph(OPENING); doc.add_paragraph(mid); doc.add_paragraph(WRAPPER); doc.add_paragraph(ASK)
    doc.add_paragraph("Best,\nAJ")
    rows=firm_info(fn)
    hp=doc.add_paragraph(); hr=hp.add_run("FIRM INFO"); hr.bold=True; hr.font.size=Pt(9); hr.font.color.rgb=RGBColor(0x80,0x80,0x80)
    if not rows:
        gp=doc.add_paragraph("(no structured firm record matched — see round-1 email)", style="List Bullet")
        for r in gp.runs: r.font.size=Pt(9); r.italic=True; r.font.color.rgb=RGBColor(0x80,0x80,0x80)
    for row in rows:
        gp=doc.add_paragraph(row, style="List Bullet")
        for r in gp.runs: r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x80,0x80,0x80)
    doc.add_paragraph("")

doc.add_page_break()
doc.add_heading("Flags for review (Jack/AJ)",level=1)
for fl in FLAGS: doc.add_paragraph(fl, style="List Bullet")

out="Newry Projects/Sequoia/Walnut Creek Second Emails - v2 DRAFT.docx"
doc.save(out)

# QA report
lo=min(wc(m) for *_,m in EMAILS); hi=max(wc(m) for *_,m in EMAILS)
avg=sum(wc(m) for *_,m in EMAILS)/len(EMAILS)
print(f"Firms drafted: {len(EMAILS)}")
print(f"Middle-paragraph words: min={lo} max={hi} avg={avg:.0f}")
over=[(fn,wc(m)) for fn,_,_,m in EMAILS if wc(m)>80]
print("Over 80w:", over if over else "none")
print("Saved:", out)
